"""
CHARAMOU AI - Contrôleur Microsoft Word
Ouvre, crée, édite et sauvegarde des documents Word.
"""
import os
import re
from pathlib import Path
from typing import Optional
from core.logger import setup_logger
from core.exceptions import AutomationError

logger = setup_logger("WordController")
DOCUMENTS_DIR = Path.home() / "Documents" / "CHARAMOU"


class WordController:
    """
    Automatise Microsoft Word via python-docx (création/édition)
    et pywin32 (contrôle COM de l'application ouverte).
    """

    def __init__(self, security=None):
        self.security = security
        DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
        logger.info("WordController initialisé.")

    def handle(self, entities: dict = None, context=None) -> str:
        """Handler pour le TaskRouter."""
        if self.security:
            self.security.require("word_control")

        entities = entities or {}
        raw_text = entities.get("raw_text", "")

        # Détecter l'action
        text_lower = raw_text.lower()
        if any(kw in text_lower for kw in ["crée", "nouveau", "créer", "ouvre un"]):
            title = self._extract_title(raw_text)
            return self.create_document(title)
        elif "ouvre" in text_lower or "ouvrir" in text_lower:
            return self.open_word()
        elif any(kw in text_lower for kw in ["écris", "rédige", "ajoute"]):
            content = self._extract_content(raw_text)
            title = self._extract_title(raw_text) or "Document CHARAMOU"
            return self.create_document(title, content=content)
        else:
            return self.open_word()

    def create_document(self, title: str = "Nouveau document", content: str = "") -> str:
        """Crée un fichier .docx avec le titre et le contenu fournis."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style titre
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contenu
            if content:
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("")

            # Sauvegarde
            safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
            filename = f"{safe_title}.docx"
            path = DOCUMENTS_DIR / filename
            doc.save(str(path))

            logger.info(f"Document créé : {path}")

            # Ouvrir le fichier
            try:
                os.startfile(str(path))
            except Exception:
                pass

            return f"Document « {title} » créé et ouvert dans Word."

        except ImportError:
            logger.error("python-docx non installé.")
            return "python-docx n'est pas installé. Exécutez : pip install python-docx"
        except Exception as e:
            logger.error(f"Erreur création document : {e}")
            return f"Erreur lors de la création du document : {e}"

    def open_word(self, filepath: str = None) -> str:
        """Ouvre Microsoft Word."""
        try:
            if filepath:
                os.startfile(filepath)
                return f"Ouverture du fichier : {Path(filepath).name}"
            else:
                os.startfile("WINWORD.EXE")
                return "Microsoft Word est ouvert."
        except Exception:
            try:
                import subprocess
                subprocess.Popen(["WINWORD.EXE"])
                return "Microsoft Word est ouvert."
            except Exception as e:
                return f"Impossible d'ouvrir Word : {e}"

    def add_text_to_open_doc(self, text: str) -> str:
        """Ajoute du texte au document Word actuellement actif (via COM)."""
        try:
            import win32com.client
            word_app = win32com.client.GetObject(None, "Word.Application")
            doc = word_app.ActiveDocument
            selection = word_app.Selection
            selection.TypeText(text + "\n")
            logger.info(f"Texte ajouté au document actif.")
            return f"Texte ajouté au document."
        except Exception as e:
            logger.warning(f"COM Word indisponible : {e}")
            return "Impossible d'accéder au document Word ouvert."

    def _extract_title(self, text: str) -> str:
        patterns = [
            r'(?:intitulé|appelé|nommé|sur|titre)\s+["\']?([^"\']+)["\']?',
            r'document\s+([A-Za-zÀ-ÿ\s]{3,30})',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "Document CHARAMOU"

    def _extract_content(self, text: str) -> str:
        match = re.search(r'(?:écris|rédige|ajoute)[:\s]+(.+)', text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return ""
