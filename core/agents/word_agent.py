"""
CHARAMOU AI - Agent Word
Création, édition et gestion intelligente de documents.
"""
import re
from pathlib import Path
from core.agents.base_agent import BaseAgent
from typing import Any


class WordAgent(BaseAgent):
    """
    Agent spécialisé dans Microsoft Word.
    Crée, édite, formate et exporte des documents.
    """

    name        = "word_agent"
    description = "Crée et édite des documents Word intelligemment."

    KEYWORDS = ["word", "document", "docx", "lettre", "rapport", "rédige", "écris", "crée un fichier"]

    def can_handle(self, task: str, entities: dict) -> bool:
        return any(kw in task.lower() for kw in self.KEYWORDS)

    def execute(self, task: str, entities: dict, context: Any = None) -> str:
        self._log_step(f"Tâche Word : '{task}'")

        task_lower = task.lower()

        if any(kw in task_lower for kw in ["crée", "nouveau", "rédige", "écris"]):
            return self._create_document(task, entities, context)
        elif "ouvre" in task_lower:
            return self._open_word()
        elif any(kw in task_lower for kw in ["lis", "résume", "analyse"]):
            return self._read_document(entities)
        else:
            return self._create_document(task, entities, context)

    def _create_document(self, task: str, entities: dict, context: Any) -> str:
        title   = self._extract_title(task)
        content = self._generate_content(task, context)

        self._log_step(f"Création : '{title}'")
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Titre
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contenu
            if content:
                for para in content.split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())

            # Date
            import datetime
            doc.add_paragraph(f"\nCréé par CHARAMOU AI — {datetime.date.today().strftime('%d/%m/%Y')}")

            # Sauvegarde
            docs_dir = Path.home() / "Documents" / "CHARAMOU"
            docs_dir.mkdir(parents=True, exist_ok=True)
            safe = re.sub(r'[\\/:*?"<>|]', '_', title)
            path = docs_dir / f"{safe}.docx"
            doc.save(str(path))

            try:
                import os
                os.startfile(str(path))
            except Exception:
                pass

            return f"Document « {title} » créé et ouvert dans Word."
        except ImportError:
            return "python-docx non installé : pip install python-docx"

    def _read_document(self, entities: dict) -> str:
        path = entities.get("file_path")
        if not path:
            return "Précisez le chemin du document à lire."
        try:
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if self.memory:
                self.memory.add_knowledge(path, text, summary=text[:100])
            return f"Document lu : {len(text)} caractères. '{text[:200]}...'"
        except Exception as e:
            return f"Erreur lecture document : {e}"

    def _open_word(self) -> str:
        try:
            import os
            os.startfile("WINWORD.EXE")
            return "Microsoft Word est ouvert."
        except Exception:
            return "Impossible d'ouvrir Word."

    def _extract_title(self, task: str) -> str:
        patterns = [r'(?:intitulé|sur|titre|appelé)\s+["\']?([^"\'\.]{3,40})',
                    r'(?:crée|rédige|écris)\s+(?:un|une|le|la)?\s+([a-zéàèê\s]{3,30})']
        for p in patterns:
            m = re.search(p, task, re.IGNORECASE)
            if m:
                return m.group(1).strip().title()
        return "Document CHARAMOU"

    def _generate_content(self, task: str, context: Any) -> str:
        # Si IA disponible, demander au modèle de générer le contenu
        if context and hasattr(context, 'get_openai_messages'):
            return ""   # Contenu vide, Word s'ouvrira vide
        return ""
